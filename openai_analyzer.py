import datetime
import json
import logging
from dataclasses import dataclass
from typing import Dict, List, Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from openai import OpenAI

from config import OpenAIConfig

logger = logging.getLogger(__name__)

@dataclass
class MeetingAnalysis:
    """Структура для анализа технического совещания"""
    transcript: str #Текст совещания
    summary: str #Сводка
    tasks: List[Dict[str, Any]] #Поставленные задачи
    hypotheses: List[Dict[str, str]] #Поставленные гипотезы
    decisions: List[str] #Принятые решения
    participants: List[str] #Участники совещания
    president: str #Председатель
    secretary: str #Секретарь
    absent: List[str] #Отсутствовавшии

class OpenAIAnalyzer:
    """Класс для анализа транскрипции с помощью OpenAI"""

    def __init__(self, config: OpenAIConfig):
        """
        Инициализация анализатора

        Args:
            config: Конфигурация OpenAI
        """
        self.config = config
        self.client = OpenAI(api_key=config.api_key)

        # Определение схемы для tool calling
        self.meeting_analysis_tool = {
            "type": "function",
            "function": {
                "name": "analyze_technical_meeting",
                "description": "Анализирует транскрипцию технического совещания и извлекает структурированную информацию",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "summary": {
                            "type": "string",
                            "description": "Краткое резюме совещания в 2-3 предложениях"
                        },
                        "tasks": {
                            "type": "array",
                            "description": "Список выявленных задач с полной информацией",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "название": {
                                        "type": "string",
                                        "description": "Название задачи (краткое и емкое)"
                                    },
                                    "описание": {
                                        "type": "string",
                                        "description": "Подробное описание задачи с техническими деталями"
                                    },
                                    "суть_задачи": {
                                        "type": "string",
                                        "description": "Краткая суть задачи в 1-2 предложениях, основная цель"
                                    },
                                    "кто_выполняет": {
                                        "type": "string",
                                        "description": "Ответственный за выполнение (имя, должность или отдел)"
                                    },
                                    "срок": {
                                        "type": "string",
                                        "description": "Срок выполнения в формате YYYY-MM-DD, в текстовом формате (завтра/послезавтра/через неделю/через две недели/через меняц) или строка Не указан, если в совещании не обговаривалось"
                                    }
                                },
                                "required": ["название", "описание", "суть_задачи", "кто_выполняет", "срок"]
                            }
                        },
                        "hypotheses": {
                            "type": "array",
                            "description": "Список гипотез, требующих проверки",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "hypothesis": {
                                        "type": "string",
                                        "description": "Описание гипотезы"
                                    },
                                    "status": {
                                        "type": "string",
                                        "enum": ["требует проверки", "принята", "отклонена"],
                                        "description": "Статус гипотезы"
                                    },
                                    "related_area": {
                                        "type": "string",
                                        "description": "Связанная техническая область"
                                    }
                                },
                                "required": ["hypothesis", "status"]
                            }
                        },
                        "decisions": {
                            "type": "array",
                            "description": "Список принятых решений",
                            "items": {
                                "type": "string",
                                "description": "Описание принятого решения"
                            }
                        },
                        "participants": {
                            "type": "array",
                            "description": "Список участников совещания",
                            "items": {
                                "type": "string",
                                "description": "Фамилия и инициалы участника"
                            }
                        },
                        "president": {
                            "type": "string",
                            "description": "Фамилия и инициалы председателя совещания"
                        },
                        "secretary": {
                            "type": "string",
                            "description": "Фамилия и инициалы секретаря совещания"
                        },
                        "absent": {
                            "type": "array",
                            "description": "Список отсутствовавших на совещании",
                            "items": {
                                "type": "string",
                                "description": "Фамилия и инициалы отсутствовавшего"
                            }
                        }

                    },
                    "required": ["summary", "tasks", "hypotheses", "decisions", "participants", "president", "secretary", "absent"]
                }
            }
        }

    def create_analysis_prompt(self, transcript: str) -> str:
        """
        Создание промпта для анализа

        Args:
            transcript: Текст транскрипции

        Returns:
            str: Промпт для анализа
        """
        return f"""
        Проанализируй транскрипцию технического совещания на предприятии.
        НЕ ДОБАВЛЯЙ НИЧЕГО ОТ СЕБЯ, ИСПОЛЬЗУЙ ТОЛЬКО ИНФОРМАЦИЮ С СОВЕЩАНИЯ, ЕСЛИ НА СОВЕЩАНИИ НЕ ХВАТИЛО ИНФОРМАЦИИ
        О ЧЕМ-ЛИБО ПОМЕЧАЙ КАК \"Не указано\"

        Для каждой задачи обязательно заполни все основные поля:
        - название: краткое название задачи
        - описание: подробное техническое описание
        - суть_задачи: краткая суть в 1-2 предложениях
        - кто_выполняет: конкретный исполнитель (если не указан, то пиши \"Не указан\")
        - срок: конкретная дата или период выполнения (если не указан, то пиши \"Не указан\")

        Обрати особое внимание на:
        - Технические решения и их обоснование
        - Проблемы, которые нужно решить
        - Распределение ответственности между участниками
        - Временные рамки выполнения задач

        Транскрипция совещания:
        {transcript}
        """

    def parse_tool_response(self, response) -> Dict[str, Any]:
        """
        Парсинг ответа от tool calling

        Args:
            response: Ответ от OpenAI

        Returns:
            Dict: Распарсенные данные
        """
        try:
            tool_call = response.choices[0].message.tool_calls[0]
            return json.loads(tool_call.function.arguments)
        except (IndexError, KeyError, json.JSONDecodeError) as e:
            logger.error(f"Ошибка парсинга ответа tool calling: {e}")
            raise

    def analyze_transcript(self, transcript: str) -> MeetingAnalysis:
        """
        Анализ транскрипции совещания

        Args:
            transcript: Текст транскрипции

        Returns:
            MeetingAnalysis: Структурированный анализ
        """
        logger.info("Начало анализа транскрипции с OpenAI...")

        if not transcript.strip():
            logger.warning("Пустая транскрипция для анализа")
            return self._create_empty_analysis(transcript)

        try:
            analysis_prompt = self.create_analysis_prompt(transcript)

            response = self.client.chat.completions.create(
                model=self.config.analyze_model,
                messages=[
                    {
                        "role": "system",
                        "content": "Ты эксперт по анализу технических совещаний на предприятиях. "
                                   "Ты специализируешься на выделении задач, гипотез и решений из "
                                   "технических дискуссий инженеров разных специальностей. "
                                   "Для транскрипции совещания использовалась модель, поддерживающая"
                                   "только русские слова, поэтому могут возникнуть ошибки, англоицизмы или технические"
                                   "термины могут быть переведены в текст, как созвучные слова,"
                                   "Обрати на это внимание."
                                   "Для каждой задачи ты ОБЯЗАТЕЛЬНО заполняешь все поля."
                                   "ЕСЛИ НЕ ХВАТАЕТ ИНФОРМАЦИИ ЗАПОЛНЯЙ ПОЛЕ КАК \"Не указан\""
                                   "НЕ ДОПОЛНЯЙ ПОЛЯ ОТ СЕБЯ, ИСПОЛЬЗУЙ ТОЛЬКО ИНФОРМАЦИЮ С СОВЕЩАНИЙ"
                    },
                    {
                        "role": "user",
                        "content": analysis_prompt
                    }
                ],
                tools=[self.meeting_analysis_tool],
                tool_choice={"type": "function", "function": {"name": "analyze_technical_meeting"}},
                temperature=self.config.temperature
            )

            analysis_data = self.parse_tool_response(response)

            logger.info(f"Анализ завершен. Найдено задач: {len(analysis_data.get('tasks', []))}")

            return MeetingAnalysis(
                transcript=transcript,
                summary=analysis_data.get("summary", ""),
                tasks=analysis_data.get("tasks", []),
                hypotheses=analysis_data.get("hypotheses", []),
                decisions=analysis_data.get("decisions", []),
                participants=analysis_data.get("participants", []),
                president=analysis_data.get("president", ""),
                secretary=analysis_data.get("secretary", ""),
                absent=analysis_data.get("absent", [])
            )

        except Exception as e:
            logger.error(f"Ошибка при анализе с OpenAI: {e}")
            return self._create_empty_analysis(transcript, str(e))

    def _create_empty_analysis(self, transcript: str, error_message: str = "") -> MeetingAnalysis:
        """
        Создание пустого анализа в случае ошибки

        Args:
            transcript: Исходная транскрипция
            error_message: Сообщение об ошибке

        Returns:
            MeetingAnalysis: Пустой анализ
        """
        summary = f"Ошибка при создании резюме: {error_message}" if error_message else "Анализ не выполнен"

        return MeetingAnalysis(
            transcript=transcript,
            summary=summary,
            tasks=[],
            hypotheses=[],
            decisions=[],
            participants=[],
            president="",
            secretary="",
            absent=[]
        )

    def validate_analysis(self, analysis: MeetingAnalysis) -> bool:
        """
        Валидация результатов анализа

        Args:
            analysis: Анализ для валидации

        Returns:
            bool: True если анализ валиден
        """
        if not analysis.summary:
            logger.warning("Отсутствует резюме совещания")
            return False

        # Проверка структуры задач
        for i, task in enumerate(analysis.tasks):
            required_fields = ["название", "описание", "суть_задачи", "кто_выполняет", "срок"]
            for field in required_fields:
                if not task.get(field):
                    logger.warning(f"Задача {i + 1}: отсутствует поле '{field}'")
                    return False

        logger.info(f"Анализ валиден: {len(analysis.tasks)} задач, {len(analysis.decisions)} решений")
        return True

    def save_analysis_to_docx(self, analysis: MeetingAnalysis, filename: str) -> None:
        """
        Сохраняет результаты анализа совещания в DOCX-файл в формате протокола.

        Args:
            analysis (MeetingAnalysis): Объект с анализом совещания.
            filename (str): Имя файла для сохранения DOCX.
        """
        document = Document()

        # Set page margins
        section = document.sections[0]
        section.top_margin = Inches(1)  # 1 inch = 2.54 cm
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        # Helper function to add a paragraph with custom formatting
        def add_formatted_paragraph(text, style_name='Normal', font_size=10, bold=False,
                                    alignment=WD_ALIGN_PARAGRAPH.LEFT, left_indent=None, space_after=None):
            paragraph = document.add_paragraph(text, style=style_name)
            run = paragraph.runs[0]
            run.font.size = Pt(font_size)
            run.bold = bold
            paragraph.alignment = alignment
            if left_indent:
                paragraph.paragraph_format.left_indent = Pt(left_indent)
            if space_after:
                paragraph.paragraph_format.space_after = Pt(space_after)
            return paragraph

        # Company Name (Жестко заданное наименование компании)
        company_name = "ОБЩЕСТВО С ОГРАНИЧЕННОЙ ОТВЕТСТВЕННОСТЬЮ НАУЧНО-ПРОИЗВОДСТВЕННОЕ ПРЕДПРИЯТИЕ \"АВТОНОМНЫЕ АЭРОКОСМИЧЕСКИЕ СИСТЕМЫ - ГЕОСЕРВИС\""
        add_formatted_paragraph(company_name, font_size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                space_after=10)

        # ПРОТОКОЛ Title
        add_formatted_paragraph("ПРОТОКОЛ", font_size=16, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=5)
        add_formatted_paragraph("технического совещания", font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                space_after=10)

        # Date and Number, City
        today = datetime.date.today()
        protocol_date = today.strftime("%d %B %Y г.").replace("январь", "января").replace("февраль", "февраля").replace(
            "март", "марта").replace("апрель", "апреля").replace("май", "мая").replace("июнь", "июня").replace("июль",
                                                                                                               "июля").replace(
            "август", "августа").replace("сентябрь", "сентября").replace("октябрь", "октября").replace("ноябрь",
                                                                                                       "ноября").replace(
            "декабрь", "декабря")
        protocol_number = f"№ {today.strftime('%Y%m%d')}-ТС"

        # Using a table for date/number and city for alignment
        header_table = document.add_table(rows=2, cols=2)
        header_table.cell(0, 0).text = protocol_date
        header_table.cell(0, 1).text = protocol_number
        header_table.cell(1, 0).text = ""  # Empty cell for alignment
        header_table.cell(1, 1).text = "г. Москва"

        # Apply formatting to table cells
        for r_idx in range(2):
            for c_idx in range(2):
                paragraph = header_table.cell(r_idx, c_idx).paragraphs[0]
                run = paragraph.runs[0]
                run.font.size = Pt(10)
                if c_idx == 1:  # Right align for number and city
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        document.add_paragraph("", style='Normal').paragraph_format.space_after = Pt(10)  # Spacer

        # Chairman and Secretary (Placeholder names)
        chairman_name = "Иванов И.И."
        secretary_name = "Петрова Е.С."
        add_formatted_paragraph(f"Председатель — {chairman_name}", font_size=10)
        add_formatted_paragraph(f"Секретарь — {secretary_name}", font_size=10, space_after=5)

        # Participants
        add_formatted_paragraph(
            f"Присутствовали: {len(analysis.participants) if analysis.participants else '0'} человек.", font_size=10)
        if analysis.participants:
            if len(analysis.participants) > 10:
                add_formatted_paragraph("(список прилагается)", font_size=10, left_indent=20)
            else:
                for p in analysis.participants:
                    add_formatted_paragraph(f"- {p}", font_size=10, left_indent=20)
        add_formatted_paragraph("Отсутствовали: 0 человек.", font_size=10,
                                space_after=10)  # Assuming 0 absent by default

        # Agenda (ПОВЕСТКА ДНЯ)
        add_formatted_paragraph("ПОВЕСТКА ДНЯ:", font_size=11, bold=True, space_after=5)
        agenda_items = []
        agenda_items.append("Обсуждение общих технических вопросов и текущего состояния проектов.")
        if analysis.tasks:
            agenda_items.append("Рассмотрение и утверждение выявленных задач.")
        if analysis.decisions:
            agenda_items.append("Принятие решений по текущим вопросам.")
        if analysis.hypotheses:
            agenda_items.append("Обсуждение и проверка гипотез.")

        for i, item in enumerate(agenda_items):
            add_formatted_paragraph(f"{i + 1}. {item}", font_size=10, left_indent=20)
        add_formatted_paragraph("").paragraph_format.space_after = Pt(10)  # Spacer

        # Main content: СЛУШАЛИ, ВЫСТУПИЛИ, РЕШИЛИ sections
        current_agenda_item_num = 1

        # Section 1: Summary of discussion
        add_formatted_paragraph(f"{current_agenda_item_num}. СЛУШАЛИ:", font_size=11, bold=True)
        add_formatted_paragraph(analysis.summary, font_size=10, left_indent=20, space_after=5)

        add_formatted_paragraph("ВЫСТУПИЛИ:", font_size=11, bold=True)
        add_formatted_paragraph("По существу обсуждения замечаний и вопросов не поступило.", font_size=10,
                                left_indent=20, space_after=5)

        add_formatted_paragraph("РЕШИЛИ:", font_size=11, bold=True)
        add_formatted_paragraph(
            f"{current_agenda_item_num}.1. Принять к сведению информацию, представленную в ходе совещания.",
            font_size=10, left_indent=20)
        add_formatted_paragraph("").paragraph_format.space_after = Pt(10)  # Spacer
        current_agenda_item_num += 1

        # Section for Tasks
        if analysis.tasks:
            add_formatted_paragraph(f"{current_agenda_item_num}. СЛУШАЛИ:", font_size=11, bold=True)
            add_formatted_paragraph("Представлены выявленные задачи, требующие выполнения в рамках проектов.",
                                    font_size=10, left_indent=20, space_after=5)

            add_formatted_paragraph("ВЫСТУПИЛИ:", font_size=11, bold=True)
            if analysis.participants:
                add_formatted_paragraph(f"{analysis.participants[0].split()[0]} – обсудил детали выполнения задач.",
                                        font_size=10, left_indent=20, space_after=5)
            else:
                add_formatted_paragraph("Участники обсудили детали каждой задачи и возможные подходы к их решению.",
                                        font_size=10, left_indent=20, space_after=5)

            add_formatted_paragraph("РЕШИЛИ:", font_size=11, bold=True)
            for i, task in enumerate(analysis.tasks):
                add_formatted_paragraph(
                    f"{current_agenda_item_num}.{i + 1}. Поручить {task.get('кто_выполняет', 'Не указан')} выполнить задачу: \"{task.get('название', 'Без названия')}\". Суть: {task.get('суть_задачи', 'Отсутствует')}. Срок: {task.get('срок', 'Не указан')}.",
                    font_size=10, left_indent=20)
                add_formatted_paragraph(f"   Подробное описание: {task.get('описание', 'Отсутствует')}", font_size=10,
                                        left_indent=40, space_after=5)
            add_formatted_paragraph("").paragraph_format.space_after = Pt(10)  # Spacer
            current_agenda_item_num += 1

        # Section for Decisions
        if analysis.decisions:
            add_formatted_paragraph(f"{current_agenda_item_num}. СЛУШАЛИ:", font_size=11, bold=True)
            add_formatted_paragraph("Представлены итоговые предложения по решениям текущих вопросов.", font_size=10,
                                    left_indent=20, space_after=5)

            add_formatted_paragraph("ВЫСТУПИЛИ:", font_size=11, bold=True)
            add_formatted_paragraph("Участники высказали свои мнения по предложенным решениям.", font_size=10,
                                    left_indent=20, space_after=5)

            add_formatted_paragraph("РЕШИЛИ:", font_size=11, bold=True)
            for i, decision in enumerate(analysis.decisions):
                add_formatted_paragraph(f"{current_agenda_item_num}.{i + 1}. {decision}", font_size=10, left_indent=20)
            add_formatted_paragraph("").paragraph_format.space_after = Pt(10)  # Spacer
            current_agenda_item_num += 1

        # Section for Hypotheses
        if analysis.hypotheses:
            add_formatted_paragraph(f"{current_agenda_item_num}. СЛУШАЛИ:", font_size=11, bold=True)
            add_formatted_paragraph("Обсуждены гипотезы, требующие дальнейшей проверки и исследования.", font_size=10,
                                    left_indent=20, space_after=5)

            add_formatted_paragraph("ВЫСТУПИЛИ:", font_size=11, bold=True)
            add_formatted_paragraph("Участники предложили методы и сроки проверки гипотез.", font_size=10,
                                    left_indent=20, space_after=5)

            add_formatted_paragraph("РЕШИЛИ:", font_size=11, bold=True)
            for i, hypothesis in enumerate(analysis.hypotheses):
                add_formatted_paragraph(
                    f"{current_agenda_item_num}.{i + 1}. Проверить гипотезу: \"{hypothesis.get('hypothesis', 'Не указана')}\". Статус: {hypothesis.get('status', 'Неизвестен')}. Связанная область: {hypothesis.get('related_area', 'Не указана')}.",
                    font_size=10, left_indent=20)
            add_formatted_paragraph("").paragraph_format.space_after = Pt(10)  # Spacer
            current_agenda_item_num += 1

        # Signatures
        document.add_paragraph("", style='Normal').paragraph_format.space_after = Pt(20)  # Spacer

        signature_table = document.add_table(rows=2, cols=2)
        signature_table.autofit = False
        # Set column widths. Inches(2) for roles, Inches(4) for signature line + name
        signature_table.columns[0].width = Inches(2)
        signature_table.columns[1].width = Inches(4)

        # Row 1: Chairman
        cell = signature_table.cell(0, 0)
        p = cell.paragraphs[0]
        p.add_run("Председательствующий").font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        cell = signature_table.cell(0, 1)
        p = cell.paragraphs[0]
        p.add_run(f"___________________ {chairman_name}").font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Row 2: Secretary
        cell = signature_table.cell(1, 0)
        p = cell.paragraphs[0]
        p.add_run("Секретарь").font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        cell = signature_table.cell(1, 1)
        p = cell.paragraphs[0]
        p.add_run(f"___________________ {secretary_name}").font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Footer
        footer_p = document.add_paragraph("", style='Normal')
        footer_p.paragraph_format.space_before = Pt(20)  # Space from signature lines
        run = footer_p.add_run("Справочник руководителя образовательного учреждения")
        run.font.size = Pt(8)
        # Note: python-docx does not directly support text color in the same way reportlab does for a simple 'grey' for the whole run.
        # You'd need to define a custom style or use more advanced techniques for exact color matching.
        # For this example, we just set font size.

        # Save the DOCX
        try:
            document.save(filename)
            logger.info(f"DOCX сохранен в файл: {filename}")
        except Exception as e:
            logger.error(f"Ошибка при создании DOCX: {e}")
            raise