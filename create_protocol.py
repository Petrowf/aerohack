from dataclasses import dataclass
from typing import Dict, List, Any
from docx import Document
import re
from copy import deepcopy
from openai_analyzer import MeetingAnalysis


def replace_placeholders(output_path: str, meeting: MeetingAnalysis, docx_path: str = "protokol_layout.docx") -> None:
    doc = Document(docx_path)

    # Регулярное выражение для плейсхолдеров: {ключ} или {модификатор:ключ}
    pattern = re.compile(r'\{(\w+)(?::(\w+))?\}')

    def get_value(key: str) -> Any:
        """Получение значения из объекта MeetingAnalysis"""
        if hasattr(meeting, key):
            return getattr(meeting, key)
        return "- данные не указаны -"

    def replace_match(match):
        modifier = match.group(1)  # Например, "count", "tableNum", "tableBig"
        key = match.group(2)  # Ключ данных (например, "tasks", "participants")

        # Обработка модификаторов
        if modifier and key:
            value = get_value(key)

            if modifier == "count" and isinstance(value, list):
                return str(len(value))

            elif modifier == "tableNum" and isinstance(value, list):
                return f"{{tableNum:{key}}}"  # Временный маркер

            elif modifier == "tableBig" and isinstance(value, list):
                return f"{{tableBig:{key}}}"  # Временный маркер

            return "- данные не указаны -"

        # Обычная замена {ключ}
        else:
            key = match.group(1)
            value = get_value(key)
            return str(value) if value is not None else "- данные не указаны -"

    # Обработка параграфов
    for paragraph in doc.paragraphs:
        full_text = ''.join(run.text for run in paragraph.runs)
        if "{" not in full_text:
            continue

        new_text = pattern.sub(replace_match, full_text)

        # Сохраняем стиль первого Run
        if paragraph.runs:
            for run in paragraph.runs[1:]:
                run.text = ""
            paragraph.runs[0].text = new_text

    # Обработка таблиц
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text = ''.join(run.text for run in paragraph.runs)
                    if "{" not in full_text:
                        continue

                    # Обработка tableNum (простые списки)
                    table_num_match = re.search(r'\{tableNum:(\w+)\}', full_text)
                    if table_num_match:
                        key = table_num_match.group(1)
                        items = get_value(key)

                        # Удаляем маркер из текущей ячейки
                        if paragraph.runs:
                            for run in paragraph.runs[1:]:
                                run.text = ""
                            paragraph.runs[0].text = ""

                        # Добавляем новые строки для каждого элемента
                        for i, item in enumerate(items, 1):
                            # Копируем строку с шаблоном
                            if i == 1:
                                # Для первого элемента используем текущую строку
                                new_row = row
                            else:
                                # Для остальных добавляем новую строку
                                new_row = deepcopy(row)
                                table._tbl.append(new_row._tr)

                            # Заменяем плейсхолдеры в новой строке
                            for new_cell in new_row.cells:
                                for new_paragraph in new_cell.paragraphs:
                                    # Сохраняем стиль
                                    if new_paragraph.runs:
                                        for run in new_paragraph.runs[1:]:
                                            run.text = ''
                                        new_paragraph.runs[0].text = str(item)
                        continue

                    # Обработка tableBig (словари)
                    table_big_match = re.search(r'\{tableBig:(\w+)\}', full_text)
                    if table_big_match:
                        key = table_big_match.group(1)
                        items = get_value(key)

                        # Очищаем ячейку с маркером
                        if paragraph.runs:
                            paragraph.runs[0].text = "1"

                        # Добавляем строки для каждого элемента
                        for i, item in enumerate(items, 1):
                            if i == 1:
                                new_row = row
                            else:
                                new_row = deepcopy(row)
                                table._tbl.append(new_row._tr)

                            # Заполняем колонки (предполагаем порядок: №, содержание, исполнитель, срок)
                            for col_idx, new_cell in enumerate(new_row.cells):
                                if col_idx == 0:
                                    new_cell.text = str(i)  # Номер
                                elif col_idx == 1:
                                    new_cell.text = item.get("суть_задачи", item.get("гипотеза", ""))
                                elif col_idx == 2:
                                    new_cell.text = item.get("кто_выполняет", item.get("ответственный", ""))
                                elif col_idx == 3:
                                    new_cell.text = item.get("срок", item.get("дата_проверки", ""))
                        continue

                    # Обычная замена плейсхолдеров в таблице
                    new_text = pattern.sub(replace_match, full_text)
                    if paragraph.runs:
                        for run in paragraph.runs[1:]:
                            run.text = ""
                        paragraph.runs[0].text = new_text

    doc.save(output_path)


if __name__ == "__main__":
    # Пример использования
    meeting_data = MeetingAnalysis(
        transcript="Полный текст совещания...",
        summary="Крутое совещание на тему мобилизации диких пчел на каторгу",
        tasks=[
            {
                "суть_задачи": "Разработать новый API",
                "кто_выполняет": "Backend-отдел",
                "срок": "15.06.2025"
            },
            {
                "суть_задачи": "Обновить дизайн",
                "кто_выполняет": "Frontend-отдел",
                "срок": "20.06.2025"
            }
        ],
        hypotheses=[
            "Пчелы могут работать без перерывов",
        ],
        decisions=["Использовать пчел для тяжелых работ", "Увеличить бюджет на мед"],
        participants=["Иван", "Мария", "Олег", "Светлана"],
        president="ROMCHIK",
        secretary="VOVCHIK",
        absent=['Вовчик', "Artem", "Sanek"]
    )

    # Запуск
    replace_placeholders("protokol_layout.docx", "protokol.docx", meeting_data)